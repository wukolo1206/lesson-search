import docx
doc = docx.Document(r'd:\test ch\閱讀教學策略查詢系統\docs\2026AI教育研討會投稿草稿.docx')
lines = []
for p in doc.paragraphs:
    if p.text.strip(): lines.append(p.text)
for t in doc.tables:
    for row in t.rows:
        row_text = ' | '.join(cell.text.replace('\n', ' ') for cell in row.cells).strip()
        if row_text: lines.append(row_text)
with open('temp_docx_content.txt', 'w', encoding='utf-8') as f:
    f.write("\n".join(lines))
